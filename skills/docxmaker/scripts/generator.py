"""
Document generator using python-docx.
Converts markdown content to Word (.docx) documents.
"""

from pathlib import Path
from typing import Optional
import re
import argparse
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_docx_from_markdown(
    content: str,
    output_file: str,
    title: Optional[str] = None
) -> str:
    """
    Convert markdown content to a Word document.
    
    Args:
        content: Markdown formatted string
        output_file: Path where the .docx file should be saved
        title: Optional document title (will be added as H1 if provided)
    
    Returns:
        Absolute path to the created document
    """
    # Create document
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title if provided
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Parse markdown and add to document
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines (but add paragraph breaks)
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Horizontal rule
        elif line.startswith('---') and not line.startswith('---|'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        # Markdown table detection
        elif line.startswith('|') and line.endswith('|'):
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one since we'll increment at end of loop
            
            # Parse table rows
            rows = []
            for tline in table_lines:
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:]+\|$', tline.replace('|', '|').replace(' ', '')):
                    continue
                # Parse cells
                cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
            
            # Create Word table if we have data
            if rows:
                num_cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = _format_inline_markdown(cell_text)
                
                # Add some space after table
                doc.add_paragraph()
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            text = _format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = _format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Regular paragraph
        else:
            text = _format_inline_markdown(line)
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    # Ensure output directory exists
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Save document
    doc.save(str(output_path))
    
    return str(output_path.absolute())


def _format_inline_markdown(text: str) -> str:
    """
    Handle inline markdown formatting (bold, italic).
    Note: python-docx doesn't easily support inline formatting,
    so we'll strip markdown syntax for now.
    TODO: Implement proper inline formatting with Run objects.
    """
    # For now, just strip markdown syntax
    # Future enhancement: parse and apply actual formatting
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # Bold+italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)       # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)           # Italic
    text = re.sub(r'`(.+?)`', r'\1', text)             # Code
    
    return text


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate Word document from markdown')
    parser.add_argument('--content', type=str, required=True, help='Markdown content')
    parser.add_argument('--output_file', type=str, required=True, help='Output .docx file path')
    parser.add_argument('--title', type=str, default=None, help='Document title')
    
    args = parser.parse_args()
    
    try:
        output = create_docx_from_markdown(
            content=args.content,
            output_file=args.output_file,
            title=args.title
        )
        print(f"✅ Document generated: {output}")
        sys.exit(0)
    except Exception as e:
        print(f"❌ Error generating document: {e}")
        sys.exit(1)
