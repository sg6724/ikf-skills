"""
Export API Route

Convert markdown content to various formats (DOCX, PDF, XLSX).
Returns downloadable files.
"""

import io
import re
from pathlib import Path
from typing import Literal

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

router = APIRouter(prefix="/api/export", tags=["export"])


class ExportRequest(BaseModel):
    """Request model for export endpoint"""
    content: str = Field(..., description="Markdown content to export")
    format: Literal["docx", "pdf", "xlsx"] = Field(..., description="Export format")
    filename: str = Field("export", description="Filename without extension")
    
    class Config:
        json_schema_extra = {
            "example": {
                "content": "# Report\n\nThis is the content...",
                "format": "docx",
                "filename": "my-report"
            }
        }


def sanitize_filename(filename: str) -> str:
    """Sanitize filename to prevent path traversal and invalid chars"""
    # Remove path separators and other dangerous chars
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    # Limit length
    return filename[:100] if filename else "export"


def markdown_to_docx(content: str) -> io.BytesIO:
    """Convert markdown to DOCX format"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=500, 
            detail="python-docx not installed. Run: uv add python-docx"
        )
    
    doc = Document()
    
    # Parse markdown and add to document
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Headers
        if line.startswith('### '):
            para = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            para = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            para = doc.add_heading(line[2:], level=1)
        
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(text, style='List Number')
        
        # Code blocks
        elif line.strip().startswith('```'):
            # Collect code block content
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        # Regular paragraphs (skip empty lines)
        elif line.strip():
            # Handle bold and italic
            text = line
            para = doc.add_paragraph()
            
            # Simple bold/italic handling
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)
        
        i += 1
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def markdown_to_xlsx(content: str) -> io.BytesIO:
    """
    Convert markdown tables to XLSX format.
    
    Extracts tables from markdown and creates Excel sheets.
    Non-table content is placed in a "Content" sheet.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="openpyxl not installed. Run: uv add openpyxl"
        )
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Content"
    
    # Style definitions
    header_font = Font(bold=True)
    header_fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    lines = content.split('\n')
    row_num = 1
    table_count = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Detect markdown table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\|\-\s:]+$', lines[i + 1]):
            table_count += 1
            
            # Create new sheet for table
            if table_count > 1 or row_num > 1:
                ws = wb.create_sheet(title=f"Table {table_count}")
            
            # Parse header
            headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center')
            
            # Skip separator line
            i += 2
            table_row = 2
            
            # Parse data rows
            while i < len(lines) and '|' in lines[i]:
                cells = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                for col, value in enumerate(cells, 1):
                    cell = ws.cell(row=table_row, column=col, value=value)
                    cell.border = thin_border
                table_row += 1
                i += 1
            
            # Auto-adjust column widths
            for col in range(1, len(headers) + 1):
                max_length = 0
                column_letter = ws.cell(row=1, column=col).column_letter
                for row in range(1, table_row):
                    cell_value = str(ws.cell(row=row, column=col).value or "")
                    max_length = max(max_length, len(cell_value))
                ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
            
            continue
        
        # Non-table content goes to first sheet
        if line.strip():
            # Reset to first sheet if we're adding non-table content after tables
            if table_count > 0 and ws.title != "Content":
                ws = wb["Content"]
            
            # Add text content
            if line.startswith('#'):
                # Header - make it bold
                level = len(re.match(r'^#+', line).group())
                text = line.lstrip('#').strip()
                cell = ws.cell(row=row_num, column=1, value=text)
                cell.font = Font(bold=True, size=14 - level)
            else:
                ws.cell(row=row_num, column=1, value=line)
            row_num += 1
        
        i += 1
    
    # Remove empty Content sheet if we only have tables
    if table_count > 0 and wb["Content"].max_row == 1 and not wb["Content"].cell(1, 1).value:
        del wb["Content"]
    
    # Save to bytes
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def markdown_to_pdf(content: str) -> io.BytesIO:
    """
    Convert markdown to PDF format using fpdf2.
    
    Uses fpdf2 which is pure Python and doesn't require system dependencies.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="fpdf2 not installed. Run: uv add fpdf2"
        )
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    
    # Set default font - use built-in font with latin-1 encoding for safety
    pdf.set_font("Helvetica", size=11)
    
    # Calculate effective width
    effective_width = pdf.w - pdf.l_margin - pdf.r_margin
    
    def safe_text(text: str) -> str:
        """Encode text safely for PDF, replacing unsupported characters."""
        # Replace common unicode characters with ASCII equivalents
        replacements = {
            '"': '"', '"': '"', ''': "'", ''': "'",
            '–': '-', '—': '-', '…': '...', '•': '*',
            '→': '->', '←': '<-', '↔': '<->',
            '✓': '[x]', '✗': '[ ]', '✔': '[x]',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Encode to latin-1, replacing unknown chars
        return text.encode('latin-1', errors='replace').decode('latin-1')
    
    lines = content.split('\n')
    
    for line in lines:
        if not line.strip():
            pdf.ln(5)
            continue
        
        # Headers
        if line.startswith('### '):
            pdf.set_font("Helvetica", 'B', 12)
            pdf.multi_cell(effective_width, 7, safe_text(line[4:]))
            pdf.set_font("Helvetica", size=11)
        elif line.startswith('## '):
            pdf.set_font("Helvetica", 'B', 14)
            pdf.multi_cell(effective_width, 8, safe_text(line[3:]))
            pdf.set_font("Helvetica", size=11)
        elif line.startswith('# '):
            pdf.set_font("Helvetica", 'B', 16)
            pdf.multi_cell(effective_width, 10, safe_text(line[2:]))
            pdf.set_font("Helvetica", size=11)
        
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = "  * " + line.strip()[2:]
            pdf.multi_cell(effective_width, 6, safe_text(text))
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            pdf.multi_cell(effective_width, 6, safe_text("  " + line.strip()))
        
        # Regular text
        else:
            # Remove markdown formatting for PDF
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)        # Code
            pdf.multi_cell(effective_width, 6, safe_text(text))
    
    # Save to bytes
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


@router.post("")
def export_content(request: ExportRequest):
    """
    Export markdown content to DOCX, PDF, or XLSX format.
    
    Returns the converted file as a download.
    
    Supported formats:
    - **docx**: Word document with basic formatting
    - **pdf**: PDF document (uses fpdf2, no system deps)
    - **xlsx**: Excel spreadsheet (tables extracted from markdown)
    """
    filename = sanitize_filename(request.filename)
    
    if request.format == "docx":
        buffer = markdown_to_docx(request.content)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    
    elif request.format == "pdf":
        buffer = markdown_to_pdf(request.content)
        media_type = "application/pdf"
        ext = "pdf"
    
    elif request.format == "xlsx":
        buffer = markdown_to_xlsx(request.content)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ext = "xlsx"
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {request.format}")
    
    return StreamingResponse(
        buffer,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}.{ext}"'
        }
    )
